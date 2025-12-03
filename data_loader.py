"""
Multi-format Data Loader for Vertex AI Gemini Integration
Supports: Images, Text, DOCX, XLSX files
"""

import os
import base64
import mimetypes
from pathlib import Path
from typing import Union, List, Dict, Any, Optional
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from enum import Enum

# Document processing
import docx
import openpyxl
from PIL import Image
import io


class DataType(Enum):
    IMAGE = "image"
    TEXT = "text"
    DOCX = "docx"
    XLSX = "xlsx"


dataclass
class LoadedData:
    """Container for loaded data with metadata"""
    content: Any
    data_type: DataType
    source_path: str
    mime_type: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None
    _raw_bytes: Optional[bytes] = field(default=None, repr=False)
    
    def get_pil_image(self) -> Optional[Image.Image]:
        """
        Get PIL Image object for display or manipulation.
        Only works for IMAGE data type.
        
        Returns:
            PIL Image object or None if not an image
        
        Example:
            >>> data = loader.load('photo.jpg')
            >>> img = data.get_pil_image()
            >>> img.show()  # Opens in default image viewer
        """
        if self.data_type != DataType.IMAGE:
            return None
        
        if self._raw_bytes:
            return Image.open(io.BytesIO(self._raw_bytes))
        
        # Decode from base64 if raw bytes not available
        image_bytes = base64.b64decode(self.content)
        return Image.open(io.BytesIO(image_bytes))
    
    def get_image_bytes(self) -> Optional[bytes]:
        """
        Get raw image bytes.
        Only works for IMAGE data type.
        
        Returns:
            Raw bytes or None if not an image
        """
        if self.data_type != DataType.IMAGE:
            return None
        
        if self._raw_bytes:
            return self._raw_bytes
        
        return base64.b64decode(self.content)
    
    def display_image(self, figsize: tuple = (10, 10), title: Optional[str] = None):
        """
        Display image using matplotlib (works in Jupyter notebooks).
        Only works for IMAGE data type.
        
        Args:
            figsize: Figure size as (width, height)
            title: Optional title for the image
        
        Example:
            >>> data = loader.load('photo.jpg')
            >>> data.display_image(figsize=(12, 8), title="My Photo")
        """
        if self.data_type != DataType.IMAGE:
            print(f"Cannot display: data type is {self.data_type.value}, not image")
            return
        
        try:
            import matplotlib.pyplot as plt
            
            img = self.get_pil_image()
            plt.figure(figsize=figsize)
            plt.imshow(img)
            plt.axis('off')
            
            if title:
                plt.title(title)
            elif self.source_path:
                plt.title(Path(self.source_path).name)
            
            plt.tight_layout()
            plt.show()
        except ImportError:
            print("matplotlib is required for display_image(). Install with: pip install matplotlib")
            # Fallback to PIL show
            img = self.get_pil_image()
            if img:
                img.show()
    
    def save_image(self, output_path: str, format: Optional[str] = None, **kwargs):
        """
        Save image to a new file.
        Only works for IMAGE data type.
        
        Args:
            output_path: Path where to save the image
            format: Image format (e.g., 'PNG', 'JPEG'). Auto-detected if None.
            **kwargs: Additional arguments passed to PIL save()
        
        Example:
            >>> data = loader.load('photo.jpg')
            >>> data.save_image('photo_copy.png')  # Converts to PNG
        """
        if self.data_type != DataType.IMAGE:
            raise ValueError(f"Cannot save: data type is {self.data_type.value}, not image")
        
        img = self.get_pil_image()
        if img:
            img.save(output_path, format=format, **kwargs)
            print(f"Image saved to: {output_path}")
    
    def resize_image(self, width: int, height: int, resample=Image.LANCZOS) -> 'LoadedData':
        """
        Resize image and return new LoadedData.
        Only works for IMAGE data type.
        
        Args:
            width: New width
            height: New height
            resample: Resampling filter (default: LANCZOS for high quality)
        
        Returns:
            New LoadedData with resized image
        
        Example:
            >>> data = loader.load('large_photo.jpg')
            >>> small = data.resize_image(800, 600)
            >>> small.display_image()
        """
        if self.data_type != DataType.IMAGE:
            raise ValueError(f"Cannot resize: data type is {self.data_type.value}, not image")
        
        img = self.get_pil_image()
        resized = img.resize((width, height), resample=resample)
        
        # Convert back to bytes and base64
        buffer = io.BytesIO()
        img_format = img.format or 'PNG'
        resized.save(buffer, format=img_format)
        new_bytes = buffer.getvalue()
        new_base64 = base64.b64encode(new_bytes).decode('utf-8')
        
        return LoadedData(
            content=new_base64,
            data_type=DataType.IMAGE,
            source_path=self.source_path,
            mime_type=self.mime_type,
            metadata={
                **self.metadata,
                'width': width,
                'height': height,
                'original_width': self.metadata.get('width'),
                'original_height': self.metadata.get('height')
            },
            _raw_bytes=new_bytes
        )
    
    def get_base64_data_uri(self) -> Optional[str]:
        """
        Get base64 data URI for embedding in HTML.
        Only works for IMAGE data type.
        
        Returns:
            Data URI string like 'data:image/png;base64,...'
        
        Example:
            >>> data = loader.load('photo.jpg')
            >>> uri = data.get_base64_data_uri()
            >>> html = f'<img src="{uri}">'
        """
        if self.data_type != DataType.IMAGE:
            return None
        
        return f"data:{self.mime_type};base64,{self.content}"
    
    def _repr_html_(self):
        """
        HTML representation for Jupyter notebooks.
        Automatically displays images inline in notebooks.
        """
        if self.data_type == DataType.IMAGE:
            uri = self.get_base64_data_uri()
            width = min(self.metadata.get('width', 500), 500)
            return f'<img src="{uri}" width="{width}"><br><small>{self.source_path}</small>'
        else:
            return f"<pre>LoadedData({self.data_type.value}): {self.source_path}</pre>"


class BaseLoader(ABC):
    """Abstract base class for data loaders"""
    
    @abstractmethod
    def load(self, file_path: str) -> LoadedData:
        pass
    
    @abstractmethod
    def supports(self, file_path: str) -> bool:
        pass


class ImageLoader(BaseLoader):
    """Loader for image files (PNG, JPG, JPEG, GIF, WEBP)"""
    
    SUPPORTED_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp'}
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def load(self, file_path: str) -> LoadedData:
        with open(file_path, 'rb') as f:
            image_bytes = f.read()
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # Get image metadata
        with Image.open(file_path) as img:
            metadata = {
                'width': img.width,
                'height': img.height,
                'format': img.format,
                'mode': img.mode
            }
        
        # Encode to base64 for Gemini API
        base64_image = base64.b64encode(image_bytes).decode('utf-8')
        
        return LoadedData(
            content=base64_image,
            data_type=DataType.IMAGE,
            source_path=file_path,
            mime_type=mime_type,
            metadata=metadata,
            _raw_bytes=image_bytes  # Store raw bytes for display
        )


class TextLoader(BaseLoader):
    """Loader for text files (TXT, MD, CSV, JSON, etc.)"""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.csv', '.json', '.xml', '.yaml', '.yml', '.log'}
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def load(self, file_path: str) -> LoadedData:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        metadata = {
            'file_size': os.path.getsize(file_path),
            'line_count': content.count('\n') + 1,
            'char_count': len(content)
        }
        
        return LoadedData(
            content=content,
            data_type=DataType.TEXT,
            source_path=file_path,
            mime_type='text/plain',
            metadata=metadata
        )


class DocxLoader(BaseLoader):
    """Loader for Microsoft Word documents"""
    
    SUPPORTED_EXTENSIONS = {'.docx'}
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def load(self, file_path: str) -> LoadedData:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs]
        full_text = '\n'.join(paragraphs)
        
        # Extract tables if present
        tables_data = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_rows.append(row_data)
            tables_data.append(table_rows)
        
        metadata = {
            'paragraph_count': len(paragraphs),
            'table_count': len(tables_data),
            'tables': tables_data
        }
        
        return LoadedData(
            content=full_text,
            data_type=DataType.DOCX,
            source_path=file_path,
            mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            metadata=metadata
        )


class XlsxLoader(BaseLoader):
    """Loader for Microsoft Excel spreadsheets"""
    
    SUPPORTED_EXTENSIONS = {'.xlsx', '.xls'}
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def load(self, file_path: str) -> LoadedData:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        sheets_data = {}
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            rows = []
            for row in sheet.iter_rows(values_only=True):
                # Filter out completely empty rows
                if any(cell is not None for cell in row):
                    rows.append(list(row))
            sheets_data[sheet_name] = rows
        
        # Convert to text representation
        text_content = []
        for sheet_name, rows in sheets_data.items():
            text_content.append(f"=== Sheet: {sheet_name} ===")
            for row in rows:
                text_content.append('\t'.join(str(cell) if cell else '' for cell in row))
        
        metadata = {
            'sheet_count': len(workbook.sheetnames),
            'sheet_names': workbook.sheetnames,
            'raw_data': sheets_data
        }
        
        return LoadedData(
            content='\n'.join(text_content),
            data_type=DataType.XLSX,
            source_path=file_path,
            mime_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            metadata=metadata
        )


class UniversalDataLoader:
    """Universal data loader that automatically selects the appropriate loader"""
    
    def __init__(self):
        self.loaders: List[BaseLoader] = [
            ImageLoader(),
            TextLoader(),
            DocxLoader(),
            XlsxLoader()
        ]
    
    def load(self, file_path: str) -> LoadedData:
        """Load a file using the appropriate loader"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        for loader in self.loaders:
            if loader.supports(file_path):
                return loader.load(file_path)
        
        raise ValueError(f"Unsupported file type: {file_path}")
    
    def load_multiple(self, file_paths: List[str]) -> List[LoadedData]:
        """Load multiple files"""
        return [self.load(fp) for fp in file_paths]
    
    def load_directory(self, directory: str, recursive: bool = False) -> List[LoadedData]:
        """Load all supported files from a directory"""
        loaded_data = []
        path = Path(directory)
        
        pattern = '**/*' if recursive else '*'
        for file_path in path.glob(pattern):
            if file_path.is_file():
                try:
                    loaded_data.append(self.load(str(file_path)))
                except ValueError:
                    # Skip unsupported files
                    continue
        
        return loaded_data